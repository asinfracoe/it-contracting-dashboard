"""
Multi-format file processor with aggressive fallbacks.
Each format tries 2-3 methods before giving up.
"""
import os
import re
from pathlib import Path
from typing import Dict


# ============================================================================
# PDF — 3 fallback methods
# ============================================================================

def extract_pdf_pymupdf(file_path: Path) -> str:
    """Method 1: PyMuPDF (fast)."""
    import fitz
    text_parts = []
    doc = None
    try:
        doc = fitz.open(str(file_path))
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                text = page.get_text("text")
                if text.strip():
                    text_parts.append(text)
            except Exception as e:
                # Skip bad pages but continue
                continue
        return "\n\n".join(text_parts)
    except Exception as e:
        return ""
    finally:
        if doc:
            try: doc.close()
            except: pass


def extract_pdf_pymupdf_blocks(file_path: Path) -> str:
    """Method 2: PyMuPDF with 'blocks' mode (handles tables better)."""
    import fitz
    text_parts = []
    doc = None
    try:
        doc = fitz.open(str(file_path))
        for page_num in range(len(doc)):
            try:
                page = doc[page_num]
                blocks = page.get_text("blocks")
                for b in blocks:
                    if len(b) > 4 and b[4].strip():
                        text_parts.append(b[4])
            except Exception:
                continue
        return "\n".join(text_parts)
    except Exception:
        return ""
    finally:
        if doc:
            try: doc.close()
            except: pass


def extract_pdf_pdfplumber(file_path: Path) -> str:
    """Method 3: pdfplumber (slower but handles complex PDFs)."""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(str(file_path)) as pdf:
            for page in pdf.pages:
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                    # Also grab tables
                    for table in (page.extract_tables() or []):
                        rows = ["\t".join(str(c) if c else "" for c in row) for row in table]
                        text_parts.append("\n".join(rows))
                except Exception:
                    continue
        return "\n\n".join(text_parts)
    except ImportError:
        return ""
    except Exception:
        return ""


def extract_pdf(file_path: Path) -> Dict:
    """Try 3 methods, return whichever extracts most text."""
    results = []
    
    for method_name, method_fn in [
        ("pymupdf", extract_pdf_pymupdf),
        ("pymupdf-blocks", extract_pdf_pymupdf_blocks),
        ("pdfplumber", extract_pdf_pdfplumber),
    ]:
        try:
            text = method_fn(file_path)
            if text and len(text.strip()) > 100:
                results.append((method_name, text))
        except Exception:
            continue
    
    if not results:
        return {"text": "", "pages": 0, "ok": False, "error": "all PDF methods failed"}
    
    # Pick the method that got the most text
    best_method, best_text = max(results, key=lambda x: len(x[1]))
    return {"text": best_text, "pages": 1, "ok": True, "method": best_method}


# ============================================================================
# DOCX
# ============================================================================

def extract_docx(file_path: Path) -> Dict:
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        tables_text = []
        for t_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables_text.append("\n".join(rows))
        
        full_text = "\n\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n=== TABLES ===\n\n" + "\n\n".join(tables_text)
        
        return {"text": full_text, "pages": 1, "ok": bool(full_text.strip())}
    except Exception as e:
        return {"text": "", "pages": 0, "ok": False, "error": str(e)}


# ============================================================================
# XLSX / XLS
# ============================================================================

def extract_xlsx(file_path: Path) -> Dict:
    try:
        import pandas as pd
        sheets = pd.read_excel(str(file_path), sheet_name=None, engine="openpyxl")
        
        text_parts = []
        for sheet_name, df in sheets.items():
            if df.empty: continue
            df = df.fillna("").astype(str)
            sheet_text = f"--- Sheet: {sheet_name} ---\n"
            sheet_text += df.to_string(index=False, max_rows=300, max_cols=25)
            text_parts.append(sheet_text)
        
        full_text = "\n\n".join(text_parts)
        return {"text": full_text, "pages": len(sheets), "ok": bool(full_text.strip())}
    except Exception as e:
        return {"text": "", "pages": 0, "ok": False, "error": str(e)}


# ============================================================================
# CSV / TXT
# ============================================================================

def extract_csv(file_path: Path) -> Dict:
    try:
        import pandas as pd
        df = pd.read_csv(str(file_path), encoding_errors="replace").fillna("").astype(str)
        return {"text": df.to_string(index=False, max_rows=500, max_cols=25), "pages": 1, "ok": True}
    except Exception as e:
        return {"text": "", "pages": 0, "ok": False, "error": str(e)}


def extract_txt(file_path: Path) -> Dict:
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        # Skip empty placeholder files
        if len(text.strip()) < 10:
            return {"text": "", "pages": 0, "ok": False, "error": "empty file"}
        return {"text": text, "pages": 1, "ok": True}
    except Exception as e:
        return {"text": "", "pages": 0, "ok": False, "error": str(e)}


# ============================================================================
# ROUTER
# ============================================================================

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".xlsx": extract_xlsx,
    ".xls": extract_xlsx,
    ".xlsb": extract_xlsx,
    ".csv": extract_csv,
    ".txt": extract_txt,
}


def process_file(file_path: Path) -> Dict:
    ext = file_path.suffix.lower()
    extractor = EXTRACTORS.get(ext)
    if not extractor:
        return {"text": "", "pages": 0, "ok": False, "error": f"unsupported: {ext}"}
    return extractor(file_path)


def is_supported(file_path: Path) -> bool:
    # Skip placeholder .txt files like "a.txt"
    if file_path.suffix.lower() == ".txt" and file_path.stat().st_size < 20:
        return False
    return file_path.suffix.lower() in EXTRACTORS
